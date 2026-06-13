import base64
import io
import logging
import os

import httpx
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from PIL import Image
import fitz  # PyMuPDF
import docx
import openpyxl

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Copilot Studio - File to Text API")

OCR_API_KEY = os.getenv("OCR_SPACE_API_KEY", "helloworld")


class FileInput(BaseModel):
    fileName: str
    contentBytes: str  # base64 encoded


class FileOutput(BaseModel):
    fileName: str
    fileType: str
    text: str


def get_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower().lstrip(".")


def ocr_via_api(image_bytes: bytes, ext: str = "png") -> str:
    b64 = base64.b64encode(image_bytes).decode()
    response = httpx.post(
        "https://api.ocr.space/parse/image",
        data={
            "apikey": OCR_API_KEY,
            "base64Image": f"data:image/{ext};base64,{b64}",
            "language": "vie",
            "isOverlayRequired": False,
        },
        timeout=30,
    )
    response.raise_for_status()
    result = response.json()
    if result.get("IsErroredOnProcessing"):
        raise Exception(result.get("ErrorMessage", ["OCR API error"])[0])
    return "\n".join(r["ParsedText"] for r in result.get("ParsedResults", [])).strip()


def extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
            else:
                pix = page.get_pixmap(dpi=200)
                ocr_text = ocr_via_api(pix.tobytes("png"), ext="png")
                text_parts.append(ocr_text)
    return "\n".join(text_parts).strip()


def extract_from_docx(file_bytes: bytes) -> str:
    f = io.BytesIO(file_bytes)
    document = docx.Document(f)

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_from_xlsx(file_bytes: bytes) -> str:
    f = io.BytesIO(file_bytes)
    wb = openpyxl.load_workbook(f, data_only=True)

    parts = []
    for sheet in wb.worksheets:
        parts.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            if row is None:
                continue
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in row_values):
                parts.append("\t".join(row_values))

    return "\n".join(parts).strip()


def extract_from_image(file_bytes: bytes, ext: str = "png") -> str:
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return ocr_via_api(buf.getvalue(), ext="png")


EXTENSION_MAP = {
    "pdf":  ("pdf",   lambda b: extract_from_pdf(b)),
    "doc":  ("word",  lambda b: extract_from_docx(b)),
    "docx": ("word",  lambda b: extract_from_docx(b)),
    "xls":  ("excel", lambda b: extract_from_xlsx(b)),
    "xlsx": ("excel", lambda b: extract_from_xlsx(b)),
    "png":  ("image", lambda b: extract_from_image(b, "png")),
    "jpg":  ("image", lambda b: extract_from_image(b, "jpg")),
    "jpeg": ("image", lambda b: extract_from_image(b, "jpeg")),
    "bmp":  ("image", lambda b: extract_from_image(b, "bmp")),
    "tiff": ("image", lambda b: extract_from_image(b, "tiff")),
    "gif":  ("image", lambda b: extract_from_image(b, "gif")),
}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/extract", response_model=FileOutput)
def extract_text(payload: FileInput):
    ext = get_extension(payload.fileName)

    if ext not in EXTENSION_MAP:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")

    file_type, extractor = EXTENSION_MAP[ext]

    try:
        file_bytes = base64.b64decode(payload.contentBytes)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 in contentBytes")

    try:
        text = extractor(file_bytes)
    except Exception as e:
        logger.exception("Extraction failed for file=%s ext=%s", payload.fileName, ext)
        raise HTTPException(status_code=500, detail=f"Failed to extract content: {str(e)}")

    return FileOutput(fileName=payload.fileName, fileType=file_type, text=text)
